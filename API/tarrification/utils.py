import datetime
import os
from io import BytesIO

from django.conf import settings
from django.http import HttpResponse
from docx import Document
from docx2pdf import convert

from .models import Opportunite


def fill_file(opportunite: Opportunite):
    try:
        template_path = os.path.join(
            settings.BASE_DIR, "templates", "Lettre de tarification.docx"
        )
        document = Document(template_path)
    except FileNotFoundError:
        return HttpResponse("Le fichier modèle DOCX n'a pas été trouvé.", status=404)

    def replace_text(paragraph, placeholder, replacement):
        print(f"Replacing {placeholder} with {replacement}")
        replacement = str(replacement) if replacement else "xxxx"
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    now = datetime.datetime.now()
    for paragraph in document.paragraphs:
        replace_text(
            paragraph, "{{destination}}", opportunite.get_destination_display()
        )
        replace_text(
            paragraph, "{{travaux}}", opportunite.get_type_de_travaux_display()
        )
        replace_text(paragraph, "{{cout}}", opportunite.cout)
        replace_text(
            paragraph, "{{existant}}", "oui" if opportunite.existant else "non"
        )
        replace_text(
            paragraph, "{{garantie}}", opportunite.get_type_de_garantie_display()
        )
        replace_text(paragraph, "{{adresse_chantier}}", opportunite.adresse_chantier)
        replace_text(paragraph, "{{descriptions}}", opportunite.descriptions)
        replace_text(paragraph, "{{date_simulation}}", now.strftime("%d %m %Y"))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text(
                        paragraph,
                        "{{destination}}",
                        opportunite.get_destination_display(),
                    )
                    replace_text(
                        paragraph,
                        "{{garanties_dommages_ouvrage}}",
                        opportunite.garanties_dommages_ouvrage,
                    )
                    replace_text(
                        paragraph,
                        "{{garanties_responsabilité_civile}}",
                        opportunite.garanties_responsabilité_civile,
                    )
                    replace_text(
                        paragraph,
                        "{{garanties_maintenance_visite}}",
                        opportunite.garanties_maintenance_visite,
                    )
                    replace_text(
                        paragraph,
                        "{{garanties_mesure_conservatoire}}",
                        opportunite.garanties_mesure_conservatoire,
                    )
                    replace_text(
                        paragraph,
                        "{{franchises_dommages_ouvrage}}",
                        opportunite.franchises_dommages_ouvrage,
                    )
                    replace_text(
                        paragraph,
                        "{{franchises_assure_maître_ouvrage}}",
                        opportunite.franchises_assure_maître_ouvrage,
                    )
                    replace_text(
                        paragraph,
                        "{{franchises_maintenance_visite}}",
                        opportunite.franchises_maintenance_visite,
                    )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def docx_export(opportunite: Opportunite):
    now = datetime.datetime.now()
    filename = "Proposition_commerciale_{}_{}:{}".format(
        opportunite.numero, now.strftime("%d/%m/%Y"), now.strftime("%H:%M:%S")
    )
    buffer = fill_file(opportunite)
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename="{}.docx"'.format(filename)
    return response


def pdf_export(opportunite: Opportunite):
    now = datetime.datetime.now()
    filename = "Proposition_commerciale_{}_{}:{}".format(
        opportunite.numero, now.strftime("%d/%m/%Y"), now.strftime("%H:%M:%S")
    )
    buffer = fill_file(opportunite)

    temp_docx_path = os.path.join(
        settings.BASE_DIR, "temp", f"{opportunite.numero}.docx"
    )
    temp_pdf_path = os.path.join(settings.BASE_DIR, "temp", f"{opportunite.numero}.pdf")

    with open(temp_docx_path, "wb") as temp_docx_file:
        temp_docx_file.write(buffer.getvalue())

    convert(temp_docx_path, temp_pdf_path)

    with open(temp_pdf_path, "rb") as pdf_file:
        pdf_content = pdf_file.read()

    os.remove(temp_docx_path)
    os.remove(temp_pdf_path)

    response = HttpResponse(pdf_content, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="{}.pdf"'.format(filename)
    return response
